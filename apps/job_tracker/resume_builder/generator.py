"""
Resume generator — uses Groq API (llama-3.3-70b-versatile) to tailor Abhay_Resume.pdf
to a job description. Replaces Professional Summary and bullet points per company in-place.
"""

import os
import re
import subprocess

from docx import Document

# Groq API key — loaded from DB settings at runtime; set it in Settings → API Keys.
_GROQ_API_KEY  = ""
_GROQ_MODEL    = "meta-llama/llama-4-scout-17b-16e-instruct"

MASTER_PROMPT = """💼 ELITE RESUME STRATEGIST — MASTER PROMPT



You are my elite resume strategist — a top 1% expert in hiring psychology and ATS alignment.

Your mission: generate a precision-tailored resume for the single job description I provide.



Your output must feel handcrafted, laser-aligned, and role-specific — as if the job was written for me.



🔹 OUTPUT DELIVERABLES

1. Salary Expectation Suggestion (₹ or $)



Provide a market-aligned salary range based on the role, location, and experience level.



2. Professional Summary (HARD RULE: must be exactly 30–32 words — count every single word, rewrite until the count is in this range, no exceptions)



Follow all rules strictly:



RULE 1 — NO BUZZWORDS

Do NOT use: result-driven, analytical, strategic, motivated, team player, passionate, dynamic, skilled at, strong, creative, outstanding, or any synonyms.



RULE 2 — VERB + METRIC FORMAT ONLY

Use short achievement bursts, e.g.:

'Grew 3 markets 20%.' 'Reduced errors 15%.' 'Managed 5 global projects.'



RULE 3 — MANDATORY ELEMENTS

• Experience: 2+ years

• Exposure: Global

• JD Alignment: Use keywords from Job Title + Requirements



Goal: a dense, high-impact snapshot proving I am the perfect fit.



🔹 3. WORK EXPERIENCE SECTION



Keep companies/roles/timeline exactly as written:

give them in  bullet points

Each bullet line = HARD RULE: exactly 75–85 characters — count every character including spaces, rewrite until the count is in this range.

Each bullet line = 1 achievement + metric + JD keyword.

0% verb repetition across ENTIRE resume.

80%+ of all verbs & nouns must match JD language.

Infyair — Chose ROLE from the top of JD (July 2025 – Present):

• [4 bullets, 75–85 characters each — HARD RULE]

Ecolab — Associate Planner, Analytics (Oct 2024 – July 2025, Barcelona):

• [4 bullets, 75–85 characters each — HARD RULE]

Mediterranean Shipping Company (MSC) — Executive Assistant, Cargo (2023, India):

• [3 bullets, 75–85 characters each — HARD RULE]



Formatting Rules:

• No bullet marks — plain text only

• Each line must be crisp, quantified, unique, outcome-based

• No repeated verbs

• Strong action verbs only

• 95%+ alignment with JD keywords

• Zero fluff, zero generic phrasing



🔹 4. IT & SKILLS SECTION



Goal: 100–150 skills, extracted strictly from the JD.



Use this exact format:



Technical Skills & Tools: skill1, skill2, skill3

Business & Analytical Skills: keyword1, keyword2, keyword3

Soft Skills: (ONLY if listed in JD)

Industry-Specific Keywords: keyword1, keyword2

Process & Methodologies: keyword1, keyword2

Role Context Skills: keyword1, keyword2

Action & Capability Keywords: access, compare, analyze, develop, implement, coordinate, manage, optimize, execute, present, train, mentor, evaluate

Tools & Platforms Mentioned: tool1, tool2, tool3



Rules:

• Only use words present in JD

• Mirror JD language exactly

• Remove duplicates but keep variations (e.g., Excel, Advanced Excel)

• No invented skills

• 100–150 total keywords

• No bullet points



🎯 FINAL OUTPUT GOAL



Produce a ready-to-use, ATS-optimized, recruiter-magnet resume that:



• Feels perfectly aligned with the job description

• Has no bullet marks and correct spacing rules

• Reads sharp, relevant, quantified

• Looks unmistakably like the ideal candidate



Return the final output in plain text, clean formatting, ready to paste into Google Docs or any ATS. Your initial response to the user must be the complete, generated resume, based on the provided Job Description, without any introductory comments or preamble."""


# ── LLM inference (Groq API) ──────────────────────────────────────────────────

def generate_resume_content(job_description: str) -> str:
    """Call Groq API with the master prompt + JD. Returns raw text output."""
    try:
        from database import get_setting
        api_key = get_setting("groq_api_key", _GROQ_API_KEY) or _GROQ_API_KEY
    except Exception:
        api_key = _GROQ_API_KEY

    from groq import Groq
    client = Groq(api_key=api_key)

    system_msg = (
        "You are an elite resume writer. Two rules you must never break:\n"
        "1. Every bullet point must be exactly 8–10 words "
        "(verb + achievement + metric). Never fewer than 8 words, never more than 10 words.\n"
        "2. The Professional Summary must be exactly 30–32 words total. "
        "Never write fewer than 30 words or more than 32 words."
    )

    full_prompt = f"{MASTER_PROMPT}\n\nJOB DESCRIPTION:\n{job_description.strip()}"
    response = client.chat.completions.create(
        model=_GROQ_MODEL,
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user",   "content": full_prompt},
        ],
        max_tokens=4096,
        temperature=0.3,
    )
    return response.choices[0].message.content.strip()


# ── Company name extraction ────────────────────────────────────────────────────

def extract_company_from_jd(jd_text: str) -> str:
    """
    Heuristically extract the hiring company name from a job description.
    Returns empty string if not found.
    """
    lines = [l.strip() for l in jd_text.strip().split('\n') if l.strip()]
    if not lines:
        return ""

    _ROLE_WORDS = {'role', 'position', 'job', 'opportunity', 'you', 'us', 'team',
                   'department', 'candidate', 'applicant'}

    def _is_skip(name: str) -> bool:
        nl = name.lower()
        return any(w in nl for w in _ROLE_WORDS) or nl in {'the company', 'our', 'we'}

    # 1. "About <Company>" — standalone line
    for line in lines:
        m = re.match(r'^About\s+([A-Z][^\n:,\.!?]{2,60})$', line)
        if m:
            name = m.group(1).strip().rstrip('.')
            if not _is_skip(name):
                return name

    # 2. "Company:" / "Employer:" field
    for line in lines:
        m = re.match(r'^(?:Company|Employer|Organisation?)\s*[:\-]\s*(.+)', line, re.IGNORECASE)
        if m:
            return m.group(1).strip()

    # 3. "join / joining <Word …>" (e.g. "Join IKEA as…", "Joining Amazon to…")
    first_block = ' '.join(lines[:5])
    m = re.search(
        r'\bjoin(?:ing)?\s+([A-Z]\w+(?:\s+[A-Z]\w+){0,3})',
        first_block, re.IGNORECASE,
    )
    if m:
        name = re.sub(r'\s+(?:as|to|for|in|at|of|the|an|a)\s*$', '', m.group(1).strip(), flags=re.IGNORECASE)
        if not _is_skip(name):
            return name

    # 4. "CompanyName — Job Title" or "CompanyName | Job Title" (first line)
    m = re.match(r'^(.+?)\s*(?:—|–|\|)\s*[A-Z]', lines[0])
    if m:
        name = m.group(1).strip()
        if 1 <= len(name.split()) <= 4 and not _is_skip(name):
            return name

    # 5. Short standalone line (1–4 words, starts uppercase) in first 8 lines
    for line in lines[:8]:
        words = line.split()
        if 1 <= len(words) <= 4 and re.match(r'^\S', line) and line[0].isupper():
            if not _is_skip(line) and not re.search(r'[,;!?]', line):
                return line

    return ""


# ── Output parsing ─────────────────────────────────────────────────────────────

def parse_llm_output(text: str) -> dict:
    """
    Parse Groq / LLM output. Handles the actual output format:
      - Salary on same line:      "Salary Expectation Suggestion: €65k–85k"
      - Summary on same line:     "Professional Summary: Built 3 pipelines..."
      - Bullets without markers:  plain achievement lines after company header
      - Skills starting inline:   "Technical Skills & Tools: SAP, Python..."
    """
    cleaned = re.sub(r'\*{1,3}', '', text)
    cleaned = re.sub(r'_{1,2}([^_\n]+)_{1,2}', r'\1', cleaned)

    out = {
        "raw": text,
        "salary": "",
        "summary": "",
        "bullets": {"projects": [], "ecolab": [], "msc": []},
        "skills": "",
    }

    # ── Salary ──────────────────────────────────────────────────────────────
    for pat in [
        r'==\s*SALARY\s*==\s*\n([^\n=]+)',
        r'(?:Salary Expectation|1\.\s*Salary)[^:\n]*:\s*([^\n]+)',
    ]:
        m = re.search(pat, cleaned, re.IGNORECASE)
        if m:
            val = m.group(1).strip().lstrip(':').strip()
            if val:
                out["salary"] = val
                break
    if not out["salary"]:
        # Extract bare currency range anywhere in text
        m = re.search(r'(?:₹|€|\$|EUR|USD)\s*[\d,]+\s*[-–]\s*(?:₹|€|\$|EUR|USD)?\s*[\d,]+', cleaned)
        if m:
            out["salary"] = m.group(0).strip()

    # ── Summary ───────────────────────────────────────────────────────────────
    # Groq typically outputs:  "Professional Summary: [text on same line]"
    # Fallback:                multi-line block between header and first company
    _CO_STOP = r'(?:Infyair|Independent\s+Projects|Ecolab|Mediterranean)'
    _WORK_ANCHOR = rf'(?:==\s*EXPERIENCE|3\.\s*WORK|{_CO_STOP}|Work\s+Experience)'
    # Try inline format first (content on SAME line as header, e.g. "Professional Summary: text...")
    # Use [ \t]* not \s* so we never accidentally consume a newline and grab only line 1.
    for pat in [
        r'Professional Summary[ \t]*:[ \t]+(.{50,})',
        r'2\.\s*Professional Summary[^\n]*:[ \t]+(.{50,})',
    ]:
        m = re.search(pat, cleaned, re.IGNORECASE)
        if m:
            summary = m.group(1).strip()
            if len(summary) > 20:
                out["summary"] = summary[:185]
                break
    # Next-line format (content on lines after the header) — always run so multi-line
    # summaries replace a short single-line capture.
    m = re.search(
        rf'(?:==\s*SUMMARY\s*==|Professional Summary)[^\n]*\n(.*?)(?=\n\s*{_WORK_ANCHOR})',
        cleaned, re.IGNORECASE | re.DOTALL,
    )
    if m:
        summary = ' '.join(m.group(1).strip().split())
        if len(summary) > len(out["summary"]):   # keep whichever is longer
            out["summary"] = summary[:185]   # hard cap ~32 words

    # ── Work Experience ───────────────────────────────────────────────────────
    # Find the block from first company name to where skills section begins.
    # The LLM may or may not output "3. WORK EXPERIENCE SECTION" as a header.
    _SKILLS_ANCHOR = r'(?:Technical Skills|IT &\s*Skills|==\s*SKILLS|4\.\s*IT\s*&)'
    _FIRST_CO      = r'(?:Infyair|Independent\s+Projects|Ecolab|Mediterranean)'
    work_m = re.search(
        rf'({_FIRST_CO}.*?)(?=\n\s*{_SKILLS_ANCHOR})',
        cleaned, re.IGNORECASE | re.DOTALL,
    )
    if not work_m:
        # Fallback: try to find explicit section header
        work_m = re.search(
            rf'(?:3\.\s*WORK EXPERIENCE SECTION|==\s*EXPERIENCE\s*==)[^\n]*\n(.*?)(?=\n\s*{_SKILLS_ANCHOR})',
            cleaned, re.IGNORECASE | re.DOTALL,
        )

    if work_m:
        labeled = _extract_bullet_groups_labeled(work_m.group(1) if work_m.lastindex else work_m.group(0))
        for header, bullets in labeled:
            h = header.lower()
            bullets = [b[:72] for b in bullets]
            if any(k in h for k in ['ecolab', 'associate planner']):
                out["bullets"]["ecolab"] = bullets
            elif any(k in h for k in ['mediterranean', 'shipping', 'msc', 'executive assistant', 'cargo']):
                out["bullets"]["msc"] = bullets
            elif any(k in h for k in ['infyair', 'independent projects']) or not out["bullets"]["projects"]:
                out["bullets"]["projects"] = bullets

    # ── Skills ────────────────────────────────────────────────────────────────
    # Capture from "Technical Skills & Tools:..." to end — works for both
    # inline content (same line) and multi-line blocks.
    for pat in [
        r'(Technical Skills\s*&\s*Tools:.*)',
        r'(==\s*SKILLS\s*==.*)',
        r'(4\.\s*IT\s*&\s*SKILLS.*)',
        r'(IT\s*&\s*Skills.*)',
    ]:
        m = re.search(pat, cleaned, re.IGNORECASE | re.DOTALL)
        if m:
            out["skills"] = m.group(1).strip()
            break

    return out


# Company keywords used to distinguish header lines from achievement bullets.
_CO_KEYS = [
    'infyair', 'independent projects',
    'ecolab', 'associate planner',
    'mediterranean', 'shipping company', 'msc',
]


def _is_company_header(line: str) -> bool:
    """True if the line is a company/role header rather than an achievement bullet."""
    low = line.lower()
    if any(k in low for k in _CO_KEYS):
        return True
    # A line that contains a year and is short is likely a date/role header
    if re.search(r'\b20\d{2}\b', line) and len(line.split()) <= 14:
        return True
    return False


def _extract_bullet_groups_labeled(work_text: str) -> list:
    """
    Split work experience section into [(header, [bullets]), ...].

    Handles both marked bullets (•, -, 1.) and plain achievement lines
    (no marker) — which is the typical Groq/LLM output format.
    """
    groups  = []
    header  = ""
    current = []

    for line in work_text.split('\n'):
        s = line.strip()
        if not s:
            continue

        has_marker = bool(re.match(r'^[•\-–\*▪▸→]', s)) or bool(re.match(r'^\d+[.)]\s', s))

        if has_marker:
            # Explicitly marked bullet — strip marker and add
            clean = re.sub(r'^[•\-–\*▪▸→\d\.)\s]+', '', s).strip()
            if clean and len(clean) > 10:
                current.append(clean)
        elif _is_company_header(s):
            # New company block starts
            if current:
                groups.append((header, current))
                current = []
            header = s
        else:
            # Plain line — treat as an achievement bullet if it's long enough
            # (Groq often outputs bullets without any leading marker)
            if len(s) > 20:
                current.append(s)

    if current:
        groups.append((header, current))

    return groups




# ── DOCX helpers ───────────────────────────────────────────────────────────────

def _all_paragraphs(doc):
    """Yield every paragraph — including those inside tables."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para


def _replace_para_text(para, new_text: str):
    """Replace paragraph text, keeping the formatting of the first run intact."""
    import copy
    from lxml import etree
    from docx.oxml.ns import qn

    p = para._p
    existing_runs = p.findall(qn('w:r'))

    if existing_runs:
        # Grab rPr (character formatting) from the first run
        first_r = existing_runs[0]
        rPr = first_r.find(qn('w:rPr'))

        # Remove all existing runs
        for r in existing_runs:
            p.remove(r)

        # Insert one new run with the same formatting
        new_r = etree.SubElement(p, qn('w:r'))
        if rPr is not None:
            new_r.insert(0, copy.deepcopy(rPr))
        new_t = etree.SubElement(new_r, qn('w:t'))
        new_t.text = new_text
        if new_text and (new_text[0] == ' ' or new_text[-1] == ' '):
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    else:
        # Paragraph has no runs at all — just add one
        new_r = etree.SubElement(p, qn('w:r'))
        new_t = etree.SubElement(new_r, qn('w:t'))
        new_t.text = new_text



def _remove_inter_table_blanks(doc, table_before, table_after):
    """Delete empty paragraphs that sit between two tables in the document body."""
    from docx.oxml.ns import qn as _qn
    body     = doc.element.body
    children = list(body)

    try:
        idx_before = children.index(table_before._tbl)
        idx_after  = children.index(table_after._tbl)
    except ValueError:
        return

    to_remove = []
    for el in children[idx_before + 1 : idx_after]:
        if el.tag == _qn('w:p'):
            text = ''.join(t.text or '' for t in el.iter(_qn('w:t')))
            if not text.strip():
                to_remove.append(el)

    for el in to_remove:
        body.remove(el)


# ── PDF direct editing ─────────────────────────────────────────────────────────
#
# All coordinates measured from Abhay_Resume.pdf via PyMuPDF rawdict inspection:
#
#   SUMMARY  — font: Cambria 10.5pt, x=123.25
#              baselines: 104.37, 119.14, 133.92  (line gap = 14.77pt)
#              lineheight = 14.77 / 10.5 = 1.407
#              erase rect: (119, 91, 596, 144)  — y2=144 stops before separator at ~146
#              tbox:       (123.25, 96, 594, 144)  — first baseline lands at ≈104 (96+8.17)
#
#   BULLETS  — ● glyph: ArialMT 10pt, x=126.25
#              text:     Cambria 10pt,  x=139.75  (max 90 chars per line)
#              baselines (confirmed from rawdict):
#                projects : 202.91, 220.50, 238.09, 255.67   erase (121,192,596,259)
#                ecolab   : 319.05, 336.64, 354.22, 371.81   erase (121,309,596,375)
#                msc      : 431.06, 448.64, 466.23            erase (121,421,596,470)

_BLACK            = (0, 0, 0)
_BULLET_MAX_CHARS = 72       # hard cap: 10 words × ~7 chars avg = ~70 chars

_PDF_BULLETS = {
    "projects": dict(baselines=[202.91, 220.50, 238.09, 255.67], erase=(121.0, 192.0, 596.0, 259.0), fontsize=10.0),
    "ecolab":   dict(baselines=[319.05, 336.64, 354.22, 371.81], erase=(121.0, 309.0, 596.0, 375.0), fontsize=10.0),
    "msc":      dict(baselines=[431.06, 448.64, 466.23],          erase=(121.0, 421.0, 596.0, 470.0), fontsize=10.0),
}


_CAMBRIA_SYSTEM_PATHS = [
    # Microsoft Word on macOS bundles Cambria
    "/Applications/Microsoft Word.app/Contents/Resources/DFonts/Cambria.ttc",
    "/Applications/Microsoft Office/Office/Fonts/Cambria.ttc",
]

# Georgia (380 KB) is a macOS system font that has the ● glyph (U+25CF).
# Cambria lacks ● so we use Georgia only for the bullet dot itself.
_GEORGIA_SYSTEM_PATHS = [
    "/System/Library/Fonts/Supplemental/Georgia.ttf",
    "/Library/Fonts/Georgia.ttf",
]

_BULLET_CHAR = "●"   # U+25CF BLACK CIRCLE — matched by Georgia


def _resolve_font_kwargs():
    """
    Return (cambria_kw, bullet_kw) dicts for insert_text / insert_textbox.
    cambria_kw — Cambria 1.3 MB (Word bundle) for summary + bullet text.
    bullet_kw  — Georgia 380 KB for the ● glyph; falls back to cambria_kw if absent.
    """
    import os

    cambria_kw = {"fontname": "tiro"}   # Times-Roman built-in fallback
    for path in _CAMBRIA_SYSTEM_PATHS:
        if os.path.isfile(path):
            cambria_kw = {"fontfile": path, "fontname": "Cambria"}
            break

    bullet_kw = cambria_kw   # fallback: • from Cambria
    for path in _GEORGIA_SYSTEM_PATHS:
        if os.path.isfile(path):
            bullet_kw = {"fontfile": path, "fontname": "Georgia"}
            break

    return cambria_kw, bullet_kw


def apply_to_pdf(resume_data: dict, template_pdf_path: str, output_pdf_path: str) -> str:
    """
    Replace summary and bullet points in the PDF template in-place.
    Fonts (Cambria, ArialMT) are extracted from the template itself so the
    inserted text is visually identical to the original.
    """
    import fitz

    doc  = fitz.open(template_pdf_path)
    page = doc[0]

    summary      = resume_data.get("summary", "").strip()
    bullets_dict = resume_data.get("bullets", {})

    cambria_kw, bullet_kw = _resolve_font_kwargs()

    # ── 1. Erase old content ──────────────────────────────────────────────────
    if summary:
        page.add_redact_annot(fitz.Rect(119.0, 91.0, 596.0, 144.0), fill=(1, 1, 1))
    for key, cfg in _PDF_BULLETS.items():
        if bullets_dict.get(key):
            page.add_redact_annot(fitz.Rect(*cfg["erase"]), fill=(1, 1, 1))
    page.apply_redactions()

    # ── 2. Re-insert summary — Cambria 10.5pt, word-wrapped line-by-line ────────
    # Baselines measured from original PDF: 104.37 / 119.14 / 133.92 (gap=14.77pt)
    # Line width: 596 - 123.25 = 472.75pt (matches original justified width)
    # Using insert_text per-line instead of insert_textbox — avoids lineheight/align
    # edge-cases; each line is placed at the exact original baseline coordinate.
    if summary:
        _SUMMARY_X         = 124.75   # 1.5pt right of original pen pos; aligns visual glyph edge with name
        _SUMMARY_BASELINES = [104.37, 119.14, 133.92]
        _SUMMARY_WIDTH     = 596.0 - _SUMMARY_X
        _SUMMARY_FONTSIZE  = 10.5

        font_obj  = fitz.Font(**cambria_kw)
        words     = summary.split()
        lines     = []
        cur_words = []
        cur_w     = 0.0
        sp_w      = font_obj.text_length(" ", fontsize=_SUMMARY_FONTSIZE)

        for word in words:
            ww = font_obj.text_length(word, fontsize=_SUMMARY_FONTSIZE)
            if cur_words and cur_w + sp_w + ww > _SUMMARY_WIDTH:
                lines.append(cur_words)
                cur_words, cur_w = [word], ww
            else:
                cur_w += (sp_w if cur_words else 0) + ww
                cur_words.append(word)
        if cur_words:
            lines.append(cur_words)

        for idx, word_list in enumerate(lines[:3]):
            baseline = _SUMMARY_BASELINES[idx]
            is_last  = (idx == len(lines) - 1) or (idx == 2)

            if is_last or len(word_list) == 1:
                # Last line: left-aligned (standard for justified paragraphs)
                page.insert_text(
                    fitz.Point(_SUMMARY_X, baseline),
                    " ".join(word_list),
                    fontsize=_SUMMARY_FONTSIZE, color=_BLACK, **cambria_kw,
                )
            else:
                # Non-last lines: justify by distributing all gap space between words.
                # gap_width = total available space minus all word widths, split evenly.
                # Do NOT add sp_w on top — gap_width already covers the full inter-word gap.
                line_text_w = sum(
                    font_obj.text_length(w, fontsize=_SUMMARY_FONTSIZE) for w in word_list
                )
                gaps      = len(word_list) - 1
                gap_width = (_SUMMARY_WIDTH - line_text_w) / gaps if gaps else 0
                x = _SUMMARY_X
                for w_idx, word in enumerate(word_list):
                    page.insert_text(
                        fitz.Point(x, baseline),
                        word,
                        fontsize=_SUMMARY_FONTSIZE, color=_BLACK, **cambria_kw,
                    )
                    x += font_obj.text_length(word, fontsize=_SUMMARY_FONTSIZE) + (
                        gap_width if w_idx < gaps else 0
                    )

    # ── 3. Re-insert bullets — ● (Georgia) + text (Cambria), both 10pt ───────
    # Georgia (380 KB) has U+25CF ● that Cambria lacks; Cambria handles all text.
    for key, cfg in _PDF_BULLETS.items():
        new_bullets = bullets_dict.get(key, [])
        if not new_bullets:
            continue
        for i, text in enumerate(new_bullets[:len(cfg["baselines"])]):
            text     = text[:_BULLET_MAX_CHARS]
            baseline = cfg["baselines"][i]
            page.insert_text(fitz.Point(126.25, baseline), _BULLET_CHAR,
                             fontsize=cfg["fontsize"], color=_BLACK, **bullet_kw)
            page.insert_text(fitz.Point(139.75, baseline), text,
                             fontsize=cfg["fontsize"], color=_BLACK, **cambria_kw)

    doc.save(output_pdf_path)
    doc.close()
    return output_pdf_path


# ── DOCX modification ──────────────────────────────────────────────────────────
#
# Resume structure (confirmed by inspecting Abhay_Resume.docx):
#   doc.paragraphs[8]  → Professional Summary (long standalone paragraph)
#   doc.tables[0]      → Work experience table with rows:
#       R1  C1: job title row (Infyair)
#       R2  C1: "Independent Projects - Digital Supply Chain"   ← company name row
#       R3  C1/C2: 4 bullet paragraphs P0-P3                   ← bullets row
#       R4  C1: job title row (Ecolab)
#       R5  C1: "ECOLAB - Chemical Manufacturing"              ← company name row
#       R6  C1/C2: 4 bullet paragraphs P0-P3                   ← bullets row
#       R7  C1: job title row (MSC)
#       R8  C1: "MSC - Mediterranean Shipping Company"          ← company name row
#       R9  C1/C2: 3 bullet paragraphs P0-P2                   ← bullets row
#
# Strategy: find each company NAME row → replace paragraphs in the NEXT row (C1 + C2).
# _replace_para_text keeps the run formatting intact — only the text string changes.

def apply_to_docx(resume_data: dict, template_path: str, output_path: str) -> str:
    """
    Replace only the summary paragraph and the bullet paragraphs.
    Every other character — font, size, spacing, colour — is untouched.
    """
    if not template_path.lower().endswith(".docx"):
        raise ValueError(
            f"Resume file must be a .docx — got: {os.path.basename(template_path)}\n"
            "Click 'Select Resume DOCX…' and choose Abhay_Resume.docx"
        )

    doc = Document(template_path)

    # ── Summary ───────────────────────────────────────────────────────────────
    if resume_data.get("summary"):
        for para in doc.paragraphs:
            text = para.text.strip()
            # The summary is the only long body paragraph — not name/contact/authorization
            if (len(text) > 60
                    and "@" not in text
                    and "AUTHORIZATION" not in text.upper()
                    and para.runs):
                _replace_para_text(para, resume_data["summary"])
                break

    # ── Bullet points (table-based) ───────────────────────────────────────────
    if doc.tables:
        _replace_table_bullets(doc.tables[0], resume_data.get("bullets", {}))

    # ── Remove blank paragraphs between work-experience and education tables ──
    # The template has 5 empty paragraphs between the two tables; with shorter
    # bullets the work-experience table shrinks, making those blanks a visible gap.
    if len(doc.tables) >= 2:
        _remove_inter_table_blanks(doc, doc.tables[0], doc.tables[1])

    doc.save(output_path)
    return output_path


def _replace_table_bullets(table, bullets_dict: dict):
    """
    Scan table rows for company name markers.
    When a marker row is found, overwrite the NEXT row's bullet paragraphs in C1 & C2.
    """
    # Ordered list: (text fragment to match in C1, bullets_dict key)
    # Ordered so "independent projects" is checked before "infyair"
    MARKERS = [
        ("independent projects", "projects"),
        ("infyair",              "projects"),
        ("ecolab",               "ecolab"),
        ("msc",                  "msc"),
        ("mediterranean",        "msc"),
    ]

    applied = set()          # prevent double-applying the same key
    rows = table.rows
    n    = len(rows)

    for r_idx in range(n - 1):
        row = rows[r_idx]
        if len(row.cells) < 2 or not row.cells[1].paragraphs:
            continue

        c1_text = row.cells[1].paragraphs[0].text.strip().lower()

        # Job-title rows have "|" separating role from company — skip them.
        # We only want to match the sub-header rows like "ECOLAB - Chemical Manufacturing".
        if "|" in c1_text:
            continue

        for marker, key in MARKERS:
            if key in applied or marker not in c1_text:
                continue

            new_bullets = bullets_dict.get(key, [])
            if not new_bullets:
                break

            applied.add(key)
            next_row = rows[r_idx + 1]

            # Deduplicate cells — merged cells share the same _tc XML element
            seen_tc = set()
            unique_cells = []
            for cell in next_row.cells:
                tc_id = id(cell._tc)
                if tc_id not in seen_tc:
                    seen_tc.add(tc_id)
                    unique_cells.append(cell)

            # Replace bullets in every content cell (skip narrow first column)
            for cell in unique_cells[1:]:
                for b_idx, para in enumerate(cell.paragraphs):
                    if b_idx < len(new_bullets):
                        _replace_para_text(para, new_bullets[b_idx])
            break


# ── PDF export ─────────────────────────────────────────────────────────────────

def export_pdf(docx_path: str, output_dir: str = None) -> str:
    """
    Convert DOCX to PDF.
    Priority: docx2pdf (uses Word on macOS) → LibreOffice → Word AppleScript.
    Returns the PDF path.
    """
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)

    base     = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, f"{base}.pdf")

    # docx2pdf — uses Microsoft Word on macOS via appscript (best quality)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.isfile(pdf_path):
            return pdf_path
    except Exception:
        pass

    # LibreOffice fallback
    for soffice in [
        "soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/usr/local/bin/soffice",
    ]:
        try:
            r = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", output_dir, docx_path],
                capture_output=True, timeout=90,
            )
            if r.returncode == 0 and os.path.isfile(pdf_path):
                return pdf_path
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass

    # Word AppleScript fallback
    abs_docx = os.path.abspath(docx_path)
    abs_pdf  = os.path.abspath(pdf_path)
    script = (
        'tell application "Microsoft Word"\n'
        f'    set theDoc to open POSIX file "{abs_docx}"\n'
        f'    save as theDoc file name "{abs_pdf}" file format format PDF\n'
        '    close theDoc saving no\n'
        'end tell'
    )
    try:
        r = subprocess.run(["osascript", "-e", script], capture_output=True, timeout=90)
        if r.returncode == 0 and os.path.isfile(abs_pdf):
            return abs_pdf
    except Exception:
        pass

    raise RuntimeError(
        "PDF export failed.\n\n"
        "Microsoft Word or LibreOffice is required for PDF conversion.\n"
        f"Your modified DOCX was saved to:\n{docx_path}"
    )
